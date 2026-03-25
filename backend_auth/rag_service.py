import os
import uuid
import logging
from typing import List, Optional, Dict, Any
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables from multiple possible locations
current_dir = os.path.dirname(__file__)
env_paths = [
    os.path.join(current_dir, ".env"),
    os.path.join(os.getcwd(), ".env"),
    ".env",
]

for env_path in env_paths:
    if os.path.exists(env_path):
        load_dotenv(env_path)
        break

from langchain_huggingface import HuggingFaceEmbeddings
from llama_index.core import SimpleDirectoryReader
from llama_index.core.schema import Document, TextNode

from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import PromptTemplate

from docx import Document as DocxDocument
import pypdf

import chromadb
from chromadb.config import Settings

from pydantic import BaseModel

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentChunk(BaseModel):
    id: str
    text: str
    document_id: str
    document_name: str
    metadata: Dict[str, Any]


class QueryRequest(BaseModel):
    query: str
    document_ids: Optional[List[str]] = None
    top_k: int = 5
    system_prompt: Optional[str] = None


class QueryResponse(BaseModel):
    answer: str
    sources: List[Dict[str, Any]]


class RAGService:
    def __init__(self, storage_dir: str = None):
        if storage_dir is None:
            base_dir = Path(__file__).parent
            storage_dir = base_dir / "rag_data"

        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)

        self._init_azure_client()
        self._init_vector_store()
        self._init_embeddings()

        existing_docs = self._get_existing_documents_count()
        logger.info(f"RAG Service initialized. Existing documents: {existing_docs}")

    def _init_azure_client(self):
        """Initialize Azure OpenAI client via Langchain"""
        api_key = os.getenv("AZURE_OPENAI_API_KEY")
        endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME", "gpt-4o-mini")

        if not api_key or api_key == "your_azure_openai_api_key_here":
            raise ValueError(
                "Azure OpenAI API key not configured. Please set AZURE_OPENAI_API_KEY in .env"
            )

        endpoint = endpoint.rstrip("/")

        # Set OPENAI_API_KEY for compatibility with AzureChatOpenAI
        os.environ["OPENAI_API_KEY"] = api_key

        api_version = os.getenv("Azure_OPENAI_API_VERSION", "2024-02-15-preview")

        self.llm = AzureChatOpenAI(
            azure_endpoint=endpoint,
            azure_deployment=deployment,
            api_key=api_key,
            api_version=api_version,
            temperature=0.7,
        )
        logger.info(
            f"Azure OpenAI client initialized with endpoint: {endpoint}, model: {deployment}"
        )

    def _init_vector_store(self):
        """Initialize Chroma vector store"""
        chroma_dir = self.storage_dir / "chroma_db"
        chroma_dir.mkdir(exist_ok=True)

        self.chroma_client = chromadb.PersistentClient(
            path=str(chroma_dir), settings=Settings(anonymized_telemetry=False)
        )

        self.collection = self.chroma_client.get_or_create_collection(
            name="rag_documents", metadata={"hnsw:space": "cosine"}
        )
        logger.info("Chroma vector store initialized")

    _cached_embedding = None
    _cached_node_parser = None

    def _init_embeddings(self):
        """Initialize HuggingFace embeddings with caching"""
        if RAGService._cached_embedding is None:
            logger.info("Loading HuggingFace embedding model...")
            RAGService._cached_embedding = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                model_kwargs={"device": "cpu"},
            )
            logger.info("HuggingFace embedding model loaded and cached")

        self.embedding = RAGService._cached_embedding

        if RAGService._cached_node_parser is None:
            from llama_index.core.node_parser import MarkdownElementNodeParser

            RAGService._cached_node_parser = MarkdownElementNodeParser(
                chunk_size=1024,
                chunk_overlap=200,
            )

        self.node_parser = RAGService._cached_node_parser
        logger.info("HuggingFace embeddings initialized (using cached model)")

    def _get_existing_documents_count(self) -> int:
        """Get count of existing documents in ChromaDB"""
        try:
            all_items = self.collection.get()
            if not all_items or not all_items.get("metadatas"):
                return 0
            unique_docs = set()
            for metadata in all_items.get("metadatas", []):
                if metadata and metadata.get("document_id"):
                    unique_docs.add(metadata.get("document_id"))
            return len(unique_docs)
        except Exception as e:
            logger.warning(f"Error getting document count: {e}")
            return 0

    def _document_exists(self, file_name: str) -> str | None:
        """Check if document already exists by filename. Returns document_id if exists."""
        try:
            all_items = self.collection.get()
            if not all_items or not all_items.get("metadatas"):
                return None

            for i, metadata in enumerate(all_items.get("metadatas", [])):
                if metadata and metadata.get("document_name") == file_name:
                    return metadata.get("document_id")
            return None
        except Exception as e:
            logger.warning(f"Error checking document existence: {e}")
            return None

    def load_document(self, file_path: str, file_content: bytes, file_name: str) -> str:
        """Load and index a document"""
        existing_doc_id = self._document_exists(file_name)
        if existing_doc_id:
            logger.info(
                f"Document {file_name} already indexed with ID: {existing_doc_id}. Skipping re-embedding."
            )
            return existing_doc_id

        document_id = str(uuid.uuid4())

        temp_dir = self.storage_dir / "temp" / document_id
        temp_dir.mkdir(parents=True, exist_ok=True)

        file_ext = Path(file_name).suffix.lower()
        file_path_full = temp_dir / file_name

        try:
            # Save the file first
            with open(file_path_full, "wb") as f:
                f.write(file_content)

            # Load documents based on file type
            documents = []
            if file_ext == ".pdf":
                documents = self._load_pdf(file_path_full, file_name)
            elif file_ext in [".docx", ".doc"]:
                documents = self._load_docx(file_path_full, file_name)
            elif file_ext in [".txt", ".md"]:
                documents = self._load_text(file_path_full, file_name)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")

            if not documents:
                raise ValueError(f"No content extracted from {file_name}")

            # Add metadata to documents
            for doc in documents:
                doc.metadata = {
                    "document_id": document_id,
                    "document_name": file_name,
                }

            # Parse into nodes with chunking
            nodes = self.node_parser.get_nodes_from_documents(documents)

            # Ensure metadata is set on nodes
            for node in nodes:
                node.metadata["document_id"] = document_id
                node.metadata["document_name"] = file_name

            self._store_in_chroma(nodes, document_id)

            logger.info(
                f"Document {file_name} indexed successfully with {len(nodes)} chunks"
            )
            return document_id

        except Exception as e:
            logger.error(f"Error indexing document: {str(e)}")
            raise

    def _load_pdf(self, file_path: Path, file_name: str) -> List[Document]:
        """Load PDF file using pypdf"""
        documents = []
        try:
            pdf_reader = pypdf.PdfReader(str(file_path))
            text_content = ""

            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    text_content += f"\n\n[Page {page_num + 1}]\n{text}"

            if text_content.strip():
                doc = Document(text=text_content, metadata={"source": file_name})
                documents.append(doc)
        except Exception as e:
            logger.error(f"Error reading PDF {file_name}: {str(e)}")
            raise

        return documents

    def _load_docx(self, file_path: Path, file_name: str) -> List[Document]:
        """Load DOCX file using python-docx"""
        documents = []
        try:
            docx_doc = DocxDocument(str(file_path))
            text_content = ""

            # Extract text from paragraphs
            for paragraph in docx_doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"

            # Extract text from tables
            for table in docx_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"

            if text_content.strip():
                doc = Document(text=text_content, metadata={"source": file_name})
                documents.append(doc)
        except Exception as e:
            logger.error(f"Error reading DOCX {file_name}: {str(e)}")
            raise

        return documents

    def _load_text(self, file_path: Path, file_name: str) -> List[Document]:
        """Load TXT or MD file"""
        documents = []
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text_content = f.read()

            if text_content.strip():
                doc = Document(text=text_content, metadata={"source": file_name})
                documents.append(doc)
        except UnicodeDecodeError:
            # Fallback to binary read if UTF-8 fails
            with open(file_path, "rb") as f:
                text_content = f.read().decode("utf-8", errors="ignore")

            if text_content.strip():
                doc = Document(text=text_content, metadata={"source": file_name})
                documents.append(doc)
        except Exception as e:
            logger.error(f"Error reading text file {file_name}: {str(e)}")
            raise

        return documents

    def _store_in_chroma(self, nodes: List[TextNode], document_id: str):
        """Store document chunks in Chroma"""
        ids = []
        documents = []
        metadatas = []

        for i, node in enumerate(nodes):
            node_id = f"{document_id}_{i}"
            ids.append(node_id)
            documents.append(node.get_text())
            metadatas.append(
                {
                    "document_id": node.metadata.get("document_id", ""),
                    "document_name": node.metadata.get("document_name", ""),
                }
            )

        # Embed all documents at once for efficiency
        embeddings = self.embedding.embed_documents(documents)

        self.collection.upsert(
            ids=ids,
            embeddings=embeddings,
            documents=documents,
            metadatas=metadatas,
        )

    def get_documents(self) -> List[Dict[str, Any]]:
        """Get list of all indexed documents"""
        unique_docs = {}

        for item in self.collection.get()["metadatas"]:
            doc_id = item.get("document_id")
            doc_name = item.get("document_name")
            if doc_id and doc_id not in unique_docs:
                unique_docs[doc_id] = {
                    "id": doc_id,
                    "name": doc_name,
                    "chunk_count": sum(
                        1
                        for m in self.collection.get()["metadatas"]
                        if m.get("document_id") == doc_id
                    ),
                }

        return list(unique_docs.values())

    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its chunks"""
        try:
            all_items = self.collection.get()

            if not all_items or not all_items.get("metadatas"):
                logger.warning(f"No items found in collection for deletion")
                return True  # Nothing to delete, consider it success

            # Find all chunk IDs that belong to this document
            ids_to_delete = []
            for i, metadata in enumerate(all_items.get("metadatas", [])):
                if metadata and metadata.get("document_id") == document_id:
                    if i < len(all_items.get("ids", [])):
                        ids_to_delete.append(all_items["ids"][i])

            logger.info(
                f"Found {len(ids_to_delete)} chunks to delete for document {document_id}"
            )

            # Delete the chunks if any were found
            if ids_to_delete:
                self.collection.delete(ids=ids_to_delete)
                logger.info(
                    f"Successfully deleted {len(ids_to_delete)} chunks for document {document_id}"
                )
            else:
                logger.warning(f"No chunks found for document {document_id}")

            return True
        except Exception as e:
            logger.error(
                f"Error deleting document {document_id}: {str(e)}", exc_info=True
            )
            return False

    def query(
        self,
        query_str: str,
        document_ids: Optional[List[str]] = None,
        top_k: int = 5,
        system_prompt: Optional[str] = None,
    ) -> QueryResponse:
        """Query the RAG system"""
        try:
            query_embedding = self.embedding.embed_query(query_str)

            where = None
            if document_ids:
                where = {"document_id": {"$in": document_ids}}

            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                where=where,
            )

            context_chunks = []
            for i, doc in enumerate(results["documents"][0]):
                metadata = results["metadatas"][0][i]
                context_chunks.append(
                    {
                        "text": doc,
                        "document_id": metadata.get("document_id", ""),
                        "document_name": metadata.get("document_name", ""),
                    }
                )

            if not context_chunks:
                return QueryResponse(
                    answer="No relevant information found in the indexed documents.",
                    sources=[],
                )

            context_text = "\n\n".join(
                [
                    f"[Source: {chunk['document_name']}]\n{chunk['text']}"
                    for chunk in context_chunks
                ]
            )

            answer = self._generate_answer(query_str, context_text, system_prompt)

            sources = [
                {
                    "document_id": chunk["document_id"],
                    "document_name": chunk["document_name"],
                    "text": chunk["text"][:200] + "..."
                    if len(chunk["text"]) > 200
                    else chunk["text"],
                }
                for chunk in context_chunks
            ]

            return QueryResponse(answer=answer, sources=sources)

        except Exception as e:
            logger.error(f"Error querying: {str(e)}")
            return QueryResponse(answer=f"Error processing query: {str(e)}", sources=[])

    def _generate_answer(
        self, query: str, context: str, system_prompt: Optional[str] = None
    ) -> str:
        """Generate answer using Azure OpenAI via Langchain"""
        if system_prompt:
            template = """{system_prompt}

Context from documents:
{context}

Question: {question}

Answer:"""
            prompt = PromptTemplate.from_template(template)
            chain = prompt | self.llm
            response = chain.invoke(
                {
                    "system_prompt": system_prompt,
                    "context": context,
                    "question": query,
                }
            )
        else:
            prompt = PromptTemplate.from_template(
                """You are a helpful AI assistant. Use the following context from documents to answer the question.

Context:
{context}

Question: {question}

Instructions:
- Answer based only on the provided context
- If the context doesn't contain enough information to answer the question, say so
- Be concise but thorough
- Cite the source documents in your answer

Answer:"""
            )
            chain = prompt | self.llm
            response = chain.invoke({"context": context, "question": query})

        return response.content.strip()


rag_service: Optional[RAGService] = None


def get_rag_service() -> RAGService:
    """Get or create the RAG service singleton"""
    global rag_service
    if rag_service is None:
        try:
            rag_service = RAGService()
        except ValueError as e:
            logger.error(f"Failed to initialize RAG service: {e}")
            raise
    return rag_service
